# Importação das bibliotecas necessárias:
# Flask: para criação da aplicação web e renderização de templates HTML;
# docx: para geração e formatação de documentos do Word;
# datetime: para manipulação de datas e horários;
# io: para trabalhar com streams em memória (no caso, o arquivo gerado).
from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from datetime import datetime
import io

# Inicialização da aplicação Flask e configuração da pasta de arquivos estáticos.
app = Flask(__name__, static_folder='static')
# Chave secreta usada para sessões; em um ambiente real (essa chave deve ser mantida em segredo hein).
app.secret_key = 'supersecretkey'


# ========== VALIDAÇÃO ==========
# Lista de campos obrigatórios que o usuário precisa preencher no formulário.
REQUIRED_FIELDS = [
    'title', 'author', 'institution', 'course', 
    'abstract', 'introducao', 'body', 'conclusao', 'references'
]


def validate_inputs(form_data):
    """
    Valida os dados enviados pelo formulário.
    Para cada campo obrigatório, verifica se foi preenchido.
    Caso algum campo esteja vazio, adiciona uma mensagem de erro.
    Retorna um dicionário com os erros ou None se não houver erros.
    """
    errors = {}
    for field in REQUIRED_FIELDS:
        # Obtém o valor do campo e remove espaços em branco nas extremidades.
        if not form_data.get(field, '').strip():
            errors[field] = 'Campo obrigatório'
    return errors if errors else None


# ========== GERADOR ABNT ==========
def configure_document_styles(doc):
    """
    Configura os estilos básicos do documento, seguindo normas de formatação ABNT.
    Define a fonte, o tamanho, o espaçamento e o recuo da primeira linha.
    """
    # Seleciona o estilo "Normal" e ajusta as propriedades de fonte e parágrafo.
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5  # Espaçamento entre linhas de 1.5
    paragraph_format.first_line_indent = Cm(1.25)  # Recuo da primeira linha (padrão ABNT)
    paragraph_format.space_after = Pt(0)  # Sem espaço extra após os parágrafos


def create_header(section, title):
    """
    Cria o cabeçalho da página.
    O cabeçalho exibe uma versão resumida do título (até 50 caracteres), alinhado à direita.
    """
    header = section.header.paragraphs[0]
    header.text = title[:50]  # Limita o título a 50 caracteres
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Alinha o cabeçalho à direita
    header.style.font.size = Pt(10)  # Define o tamanho da fonte do cabeçalho


def create_cover(doc, data):
    """
    Cria a capa do documento, centralizando as informações principais:
    - Título (em letras maiúsculas, destacado e com tamanho maior)
    - Autor (em tamanho um pouco menor)
    - Instituição, curso e a data atual (formatada de forma amigável)
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centraliza o conteúdo da capa
    
    # Adiciona o título do trabalho em destaque
    title = p.add_run(f"\n\n{data['title'].upper()}\n")
    title.font.size = Pt(16)
    title.bold = True  # Título em negrito para maior destaque
    
    # Adiciona o nome do autor
    p.add_run(f"\n{data['author']}\n").font.size = Pt(14)
    
    # Adiciona informações da instituição, curso e a data atual formatada (ex: "Março de 2025")
    institution = p.add_run(
        f"\n\n{data['institution']}\n"
        f"{data['course']}\n"
        f"{datetime.now().strftime('%B de %Y')}"
    )
    institution.font.size = Pt(12)


def format_section(doc, title, content, is_main_text=True):
    """
    Formata uma seção do documento, composta por um título e o respectivo conteúdo.
    - O título é formatado sem recuo e em negrito.
    - O conteúdo é dividido em parágrafos, alinhado de forma justificada.
    Se 'is_main_text' for True, aplica recuo na primeira linha de cada parágrafo.
    """
    # Adiciona e formata o título da seção (por exemplo: INTRODUÇÃO, CONCLUSÃO, etc.)
    title_para = doc.add_paragraph(title)
    title_para.style = 'Heading2'
    title_para.bold = True
    title_para.paragraph_format.first_line_indent = Cm(0)  # Sem recuo para o título
    
    # Divide o conteúdo em parágrafos onde há duas quebras de linha e formata cada um
    for paragraph in content.split('\n\n'):
        if paragraph.strip():  # Garante que o parágrafo não esteja vazio
            p = doc.add_paragraph(paragraph.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Alinha o parágrafo de forma justificada
            if is_main_text:
                p.paragraph_format.first_line_indent = Cm(1.25)


def create_abnt_document(form_data):
    """
    Função principal que cria o documento formatado segundo as normas ABNT.
    Utiliza os dados do formulário para gerar cada seção do documento, incluindo capa, 
    seções principais (resumo, introdução, desenvolvimento e conclusão) e referências.
    Retorna o objeto do documento criado.
    """
    doc = Document()  # Cria um novo documento
    
    
    # Configurações da página (tamanho A4 e margens definidas)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT  # Orientação vertical
    section.page_width = Cm(21)   # Largura da página (A4)
    section.page_height = Cm(29.7)  # Altura da página (A4)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    

    # Aplica os estilos do documento e adiciona o cabeçalho e a capa
    configure_document_styles(doc)
    create_header(section, form_data.get('short_title', form_data['title']))
    create_cover(doc, form_data)
    doc.add_page_break()  # Insere uma quebra de página após a capa
    

    # Define as seções principais do trabalho: Resumo, Introdução, Desenvolvimento e Conclusão.
    # No caso do Resumo, não é aplicado recuo no texto.
    sections = [
        ("RESUMO", form_data['abstract'], False),  # False indica que não haverá recuo
        ("INTRODUÇÃO", form_data['introducao']),
        ("DESENVOLVIMENTO", form_data['body']),
        ("CONCLUSÃO", form_data['conclusao'])
    ]
    

    # Para cada seção, formata e adiciona o conteúdo seguido de uma quebra de página.
    for title, content, *flags in sections:
        format_section(doc, title, content, is_main_text=(False if flags else True))
        doc.add_page_break()

    
    # Adiciona a seção de referências com formatação especial
    doc.add_paragraph("REFERÊNCIAS").style = 'Heading1'
    for ref in form_data['references'].split('\n'):
        if ref.strip():
            p = doc.add_paragraph()
            # Define indentação: recuo positivo para as linhas seguintes e recuo negativo para a primeira linha
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-1.25)
            p.add_run(ref.strip()).font.name = 'Arial'
    

    # Adiciona uma página final em branco para manter a formatação adequada
    doc.add_page_break()
    last_para = doc.add_paragraph()
    last_para.add_run("\n").font.size = Pt(12)
    
    return doc  # Retorna o documento formatado


# ========== ROTAS ==========
@app.route('/', methods=['GET', 'POST'])
def index():
    """
    Rota principal da aplicação.
    - No método GET: renderiza a página inicial com o formulário.
    - No método POST: processa os dados enviados pelo formulário, valida os campos,
      gera o documento formatado e o envia para download.
    """
    if request.method == 'POST':
        # Converte os dados do formulário para um dicionário
        form_data = request.form.to_dict()
        # Valida os dados recebidos
        errors = validate_inputs(form_data)
        
        if errors:
            # Se houver erros, re-renderiza o template com as mensagens de erro e os dados já preenchidos
            return render_template('index.html', errors=errors, form_data=form_data)
        
        try:
            # Cria um buffer em memória para armazenar o documento gerado
            buffer = io.BytesIO()
            create_abnt_document(form_data).save(buffer)
            buffer.seek(0)  # Volta o ponteiro do buffer para o início, pronto para leitura
            
            # Envia o arquivo gerado para o usuário, forçando o download.
            # O nome do arquivo é dinâmico e baseado na data e hora atuais.
            return send_file(
                buffer,
                as_attachment=True,
                download_name=f"Documento_ABNT_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except Exception as e:
            # Em caso de exceção, exibe a mensagem de erro na página, mantendo os dados preenchidos
            return render_template('index.html', error=str(e), form_data=form_data)
    
    # Se o método for GET, simplesmente renderiza o formulário inicial
    return render_template('index.html')


# Ponto de entrada da aplicação: inicia o servidor Flask em modo de debug
if __name__ == '__main__':
    app.run(debug=True)
